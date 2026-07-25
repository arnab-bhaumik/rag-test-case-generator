from types import SimpleNamespace

from docx import Document
from docx.shared import RGBColor

from src.ingestion.doc_parser import _is_reddish, _style_name, detect_red_text


def test_style_name_returns_name_when_style_present():
    para = SimpleNamespace(style=SimpleNamespace(name="Heading 1"))
    assert _style_name(para) == "Heading 1"


def test_style_name_returns_empty_string_when_style_is_none():
    # Regression: python-docx returns None for para.style when the paragraph
    # references a style ID missing from the document's style catalog —
    # common in DOCX files exported from Google Docs or edited by non-Word
    # tools. This used to crash parse_docx with
    # "'NoneType' object has no attribute 'name'" on real user uploads.
    para = SimpleNamespace(style=None)
    assert _style_name(para) == ""


def test_is_reddish_pure_red():
    assert _is_reddish(RGBColor(0xFF, 0x00, 0x00)) is True


def test_is_reddish_word_default_dark_red():
    assert _is_reddish(RGBColor(0xC0, 0x00, 0x00)) is True


def test_is_reddish_black_is_not_red():
    assert _is_reddish(RGBColor(0x00, 0x00, 0x00)) is False


def test_is_reddish_blue_is_not_red():
    assert _is_reddish(RGBColor(0x00, 0x00, 0xFF)) is False


def test_is_reddish_orange_is_not_red():
    # Orange has a high green channel — should not be caught by the red-only heuristic.
    assert _is_reddish(RGBColor(0xFF, 0xA5, 0x00)) is False


def test_is_reddish_none_is_not_red():
    assert _is_reddish(None) is False


def test_detect_red_text_extracts_only_red_runs(tmp_path):
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("This part is unchanged. ")
    red_run = para.add_run("The retry limit becomes 5 attempts.")
    red_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    para.add_run(" More unchanged text.")

    doc.add_paragraph("An entirely unrelated, unmarked paragraph.")

    path = tmp_path / "redline.docx"
    doc.save(path)

    detected = detect_red_text(path)

    assert detected == "The retry limit becomes 5 attempts."


def test_detect_red_text_no_red_text_returns_empty_string(tmp_path):
    doc = Document()
    doc.add_paragraph("Nothing here is marked as changed.")
    path = tmp_path / "plain.docx"
    doc.save(path)

    assert detect_red_text(path) == ""


def test_detect_red_text_multiple_paragraphs_joined_by_newline(tmp_path):
    doc = Document()
    p1 = doc.add_paragraph()
    r1 = p1.add_run("First change.")
    r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Second change.")
    r2.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    path = tmp_path / "two_changes.docx"
    doc.save(path)

    assert detect_red_text(path) == "First change.\nSecond change."


def test_detect_red_text_non_docx_returns_empty_string(tmp_path):
    path = tmp_path / "not_a_docx.pdf"
    path.write_bytes(b"%PDF-1.4 fake content")
    assert detect_red_text(path) == ""
