"""Parses PDF/DOCX design docs into structured sections (heading + text),
which chunker.py then splits into embedding-ready chunks.

Uses pdfplumber + python-docx directly rather than `unstructured` — simpler,
no heavy extra dependencies, and sufficient for text-based docs. Revisit
`unstructured` only if OCR (scanned PDFs) or complex table extraction is
actually needed (see plan.md §3).
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document


@dataclass
class DocSection:
    heading: str | None
    text: str
    order: int
    page: int | None = None


def parse_pdf(path: str | Path) -> list[DocSection]:
    """One section per page — pdfplumber doesn't expose heading structure."""
    sections = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = (page.extract_text() or "").strip()
            if text:
                sections.append(DocSection(heading=None, text=text, order=i, page=i + 1))
    return sections


def _style_name(para) -> str:
    """python-docx returns `None` for `para.style` when the paragraph
    references a style ID that isn't in the document's style catalog —
    common in DOCX files exported from Google Docs or edited by non-Word
    tools, not just malformed ones. Treat that as body text, not a heading."""
    return para.style.name if para.style else ""


def parse_docx(path: str | Path) -> list[DocSection]:
    """One section per Heading-styled paragraph, accumulating body text until the next heading."""
    doc = Document(str(path))
    sections: list[DocSection] = []
    heading: str | None = None
    buffer: list[str] = []
    order = 0

    def flush():
        nonlocal buffer, order
        text = "\n".join(buffer).strip()
        if text:
            sections.append(DocSection(heading=heading, text=text, order=order))
            order += 1
        buffer = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if _style_name(para).startswith("Heading"):
            flush()
            heading = text
        else:
            buffer.append(text)
    flush()
    return sections


def parse_document(path: str | Path) -> list[DocSection]:
    """Dispatches to parse_pdf/parse_docx based on file extension."""
    path = Path(path)
    if path.suffix.lower() == ".pdf":
        return parse_pdf(path)
    if path.suffix.lower() == ".docx":
        return parse_docx(path)
    raise ValueError(f"Unsupported file type: {path.suffix} (expected .pdf or .docx)")


def _is_reddish(rgb) -> bool:
    """Tolerance check, not an exact-hex match — teams pick different reds
    (pure FF0000, Word's default "Dark Red" C00000, etc.). Red-dominant with
    green/blue both low; excludes orange/pink/maroon-adjacent false positives."""
    if rgb is None:
        return False
    r, g, b = rgb[0], rgb[1], rgb[2]
    return r >= 150 and g <= 80 and b <= 80


def detect_red_text(path: str | Path) -> str:
    """Scans a DOCX for runs with red font color — teams commonly mark
    new/changed requirements this way when a design doc gets reused across
    multiple releases. Returns the detected red text joined into one string
    (for pre-filling the Generate screen's scope box), or "" if none found.
    Only explicit RGB color is checked — a run colored via a Word theme
    reference without a literal RGB value won't be caught (see plan.md)."""
    path = Path(path)
    if path.suffix.lower() != ".docx":
        return ""

    doc = Document(str(path))
    found: list[str] = []
    for para in doc.paragraphs:
        red_runs = [run.text.strip() for run in para.runs if run.text.strip() and _is_reddish(run.font.color.rgb if run.font.color else None)]
        if red_runs:
            found.append(" ".join(red_runs))
    return "\n".join(found)
